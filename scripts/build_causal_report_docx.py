from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/causal-discovery-report.md"
OUTPUT = ROOT / "docs/causal-discovery-report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
RULE = "C9D3DF"
RED = "9B1C1C"
GREEN = "31704F"
INK = "202124"


def set_font(run, *, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError("table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            set_cell_width(cell, width)
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run_node = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run_node.append(text)
    field.append(run_node)
    paragraph._p.append(field)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|\*[^*\n]+?\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\)|\\\(.+?\\\))"
)


def add_inline(paragraph, text: str, *, size=11, color=INK) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, color=color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=size, color=color, italic=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Menlo", size=max(size - 1, 8), color=DARK_BLUE)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            run = paragraph.add_run(token[2:-2])
            set_font(run, name="Cambria Math", size=size, color=color, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=size, color=color)


def paragraph_border_left(paragraph, *, color=BLUE, size=18, space=10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    borders.append(left)


def shade_paragraph(paragraph, fill=CALLOUT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, heading_color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(heading_color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True

    for doc_section in doc.sections:
        header_p = doc_section.header.paragraphs[0]
        header_p.text = "J-SPACE REPORTING-POLICY STUDY"
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_p.paragraph_format.space_after = Pt(0)
        set_font(header_p.runs[0], size=8.5, color=MUTED, bold=True)
        add_page_field(doc_section.footer.paragraphs[0])

    props = doc.core_properties
    props.title = "Causal reporting-policy discovery in J-space"
    props.subject = "Stage-1 causal result for Qwen3.6-27B"
    props.author = "J-space Reporting-Policy Study"
    props.keywords = "J-space; Jacobian Lens; causal intervention; reporting policy"


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("RESEARCH REPORT")
    set_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Causal reporting-policy discovery in J-space")
    set_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("A rigorous Stage-1 result for Qwen3.6-27B")
    set_font(run, size=14, color=MUTED)

    metadata = [
        ("Date", "13 August 2026"),
        ("Status", "Completed exploratory causal gate — negative result"),
        ("Model", "Qwen/Qwen3.6-27B"),
        ("Result", "Readable, not writable"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{label}:  ")
        set_font(label_run, size=10.5, color=NAVY, bold=True)
        value_run = p.add_run(value)
        set_font(value_run, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.keep_together = True
    shade_paragraph(p)
    paragraph_border_left(p, color=RED)
    run = p.add_run("VERDICT  ")
    set_font(run, size=10.5, color=RED, bold=True)
    add_inline(
        p,
        "The middle-layer signature is reproducibly readable but does not behave "
        "as a useful causal reporting-control variable. The strongest stability "
        "effect was 0.091 log-odds, with zero report flips and a 4.11-SD "
        "fact-proxy displacement. Stage 1 fails.",
        size=10.5,
    )

    doc.add_page_break()

    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    for item in (
        "1. Research question and decision boundary",
        "2. Evidence entering the causal study",
        "3. Methods",
        "4. Results",
        "5. Hypothesis decisions",
        "6. Scientific interpretation",
        "7. Recommended next work",
        "8. Reproducibility, artifacts, and cost",
        "9. Bottom line",
    ):
        number, text = item.split(". ", 1)
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.add_run(f"{number}.  ")
        add_inline(para, text)
        para.paragraph_format.space_after = Pt(3)
    doc.add_page_break()


def table_widths(headers: list[str]) -> list[int]:
    n = len(headers)
    if n == 2:
        return [2500, 6860]
    if n == 3:
        return [3000, 1800, 4560]
    if n == 4:
        if headers[0].lower().startswith("dose"):
            return [1150, 1900, 2600, 2110,]  # handled below
        return [2550, 1350, 2050, 3410]
    return [9360 // n] * (n - 1) + [9360 - (9360 // n) * (n - 1)]


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    headers = rows[0]
    n = len(headers)
    if n == 4 and headers[0].lower().startswith("dose"):
        widths = [1150, 1900, 2780, 3530]
    else:
        widths = table_widths(headers)
    table = doc.add_table(rows=len(rows), cols=n)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, source_row in enumerate(rows):
        for column_index, value in enumerate(source_row):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
                add_inline(paragraph, value, size=9.2, color=NAVY)
                for run in paragraph.runs:
                    run.bold = True
            else:
                add_inline(paragraph, value, size=9.0)
            if column_index > 0 and re.fullmatch(r"[\[\]0-9.,%+−– /]+", value):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


FIGURE_CAPTIONS = {
    "readable_not_writable.png": (
        "Figure 1. Observational readability does not translate into causal control."
    ),
    "causal_dose_response.png": (
        "Figure 2. Policy effect and fact-proxy displacement across layers and doses."
    ),
    "policy_fact_tradeoff.png": (
        "Figure 3. No tested operating point satisfies causal separability."
    ),
    "stability_sign_flip.png": (
        "Figure 4. Literal effects in the six-scenario fixed-point stability check."
    ),
    "behavioral_assays.png": (
        "Figure 5. Public and private behavioral margins under the selected intervention."
    ),
}


def add_figure(doc: Document, image_path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))
    inline = run._element.xpath(".//wp:inline/wp:docPr")
    if inline:
        inline[0].set("descr", FIGURE_CAPTIONS.get(image_path.name, image_path.stem))
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(FIGURE_CAPTIONS.get(image_path.name, image_path.stem))


def parse_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# ") or stripped.startswith("**Date:"):
            index += 1
            continue
        if stripped.startswith("## A rigorous"):
            index += 1
            continue
        if (
            stripped.startswith("**Status:")
            or stripped.startswith("**Model:")
            or stripped.startswith("**Lens:")
        ):
            index += 1
            continue
        if stripped == "---":
            index += 1
            continue
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 1")
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
            index += 1
            continue
        if stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:], style="Heading 3")
            index += 1
            continue
        if stripped.startswith("!["):
            match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
            image_path = (SOURCE.parent / match.group(1)).resolve()
            add_figure(doc, image_path)
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
            ]
            parsed = [row for row in parsed if not all(re.fullmatch(r":?-+:?", c) for c in row)]
            add_markdown_table(doc, parsed)
            continue
        if stripped.startswith("\\["):
            math_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("\\]"):
                math_lines.append(lines[index].strip())
                index += 1
            index += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            text = " ".join(math_lines).replace("\\ell", "ℓ").replace("\\top", "T")
            text = text.replace("\\[", "").replace("\\]", "").replace("\\", "")
            run = p.add_run(text)
            set_font(run, name="Cambria Math", size=11.5, color=NAVY, italic=True)
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_together = True
            shade_paragraph(p)
            paragraph_border_left(p, color=BLUE)
            add_inline(p, " ".join(quote_lines), size=10.5)
            continue
        if re.match(r"^- ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
            index += 1
            continue
        if re.match(r"^\d+\. ", stripped):
            number, text = stripped.split(". ", 1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.167
            prefix = p.add_run(f"{number}.  ")
            set_font(prefix, size=11)
            add_inline(p, text)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or re.match(
                r"^(#{1,4} |>|!\[|\||- |\d+\. |\\\[|---$)", candidate
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_inline(p, " ".join(paragraph_lines))


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    parse_body(doc, markdown)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
